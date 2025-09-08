from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.units import inch
from reportlab.lib.colors import black, white
from docx import Document
from docx.shared import Inches
from io import BytesIO
import json
import os
from PIL import Image as PILImage
from django.core.files.base import ContentFile
from django.conf import settings


class TemplateGenerator:
    """Utility class for generating documents from templates."""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    def generate_pdf_from_template(self, template, field_data, output_path=None):
        """Generate PDF document from template and field data."""
        if not output_path:
            output_path = f"temp_document_{template.id}.pdf"
        
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        
        # Add template content based on structure
        content_structure = template.content_structure
        
        # Process each section in the template
        for section in content_structure.get('sections', []):
            story.extend(self._process_section(section, field_data))
        
        # Build PDF
        doc.build(story)
        return output_path
    
    def generate_docx_from_template(self, template, field_data, output_path=None):
        """Generate DOCX document from template and field data."""
        if not output_path:
            output_path = f"temp_document_{template.id}.docx"
        
        doc = Document()
        
        # Process template fields
        for field in template.fields.all():
            field_value = field_data.get(field.field_name, field.default_value)
            
            if field.field_type == 'text':
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{field.field_label}: {field_value}")
                run.font.name = field.font_family
                run.font.size = field.font_size
            
            elif field.field_type == 'textarea':
                doc.add_paragraph(f"{field.field_label}:")
                doc.add_paragraph(field_value)
            
            elif field.field_type == 'table':
                if isinstance(field_value, list):
                    table = doc.add_table(rows=len(field_value), cols=len(field_value[0]) if field_value else 1)
                    for i, row_data in enumerate(field_value):
                        for j, cell_data in enumerate(row_data):
                            table.cell(i, j).text = str(cell_data)
            
            elif field.field_type == 'image' and field_value:
                try:
                    doc.add_picture(field_value, width=Inches(field.width/100))
                except:
                    doc.add_paragraph(f"[Image: {field_value}]")
        
        doc.save(output_path)
        return output_path
    
    def _process_section(self, section, field_data):
        """Process a template section and return story elements."""
        story = []
        
        section_type = section.get('type', 'paragraph')
        content = section.get('content', '')
        
        # Replace placeholders with actual data
        for field_name, field_value in field_data.items():
            placeholder = f"{{{{{field_name}}}}}"
            content = content.replace(placeholder, str(field_value))
        
        if section_type == 'title':
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=self.styles['Title'],
                fontSize=18,
                spaceAfter=20
            )
            story.append(Paragraph(content, title_style))
        
        elif section_type == 'paragraph':
            story.append(Paragraph(content, self.styles['Normal']))
            story.append(Spacer(1, 12))
        
        elif section_type == 'header':
            story.append(Paragraph(content, self.styles['Heading1']))
            story.append(Spacer(1, 12))
        
        return story


class TemplateValidator:
    """Validate template structure and field data."""
    
    @staticmethod
    def validate_template_structure(structure):
        """Validate template content structure."""
        if not isinstance(structure, dict):
            return False, "Template structure must be a dictionary"
        
        required_keys = ['sections']
        for key in required_keys:
            if key not in structure:
                return False, f"Missing required key: {key}"
        
        if not isinstance(structure['sections'], list):
            return False, "Sections must be a list"
        
        return True, "Valid structure"
    
    @staticmethod
    def validate_field_data(template, field_data):
        """Validate field data against template requirements."""
        errors = []
        
        for field in template.fields.all():
            if field.is_required and field.field_name not in field_data:
                errors.append(f"Required field '{field.field_label}' is missing")
            
            if field.field_name in field_data:
                value = field_data[field.field_name]
                
                if field.field_type == 'number':
                    try:
                        float(value)
                    except ValueError:
                        errors.append(f"Field '{field.field_label}' must be a number")
                
                elif field.field_type == 'date':
                    # Add date validation if needed
                    pass
        
        return len(errors) == 0, errors


def create_default_templates():
    """Create some default templates for common document types."""
    from .models import DocumentTemplate, TemplateField
    
    # Business Letter Template
    letter_template = DocumentTemplate.objects.create(
        name="Professional Business Letter",
        description="Standard business letter template",
        template_type="letter",
        content_structure={
            "sections": [
                {"type": "header", "content": "{{company_name}}"},
                {"type": "paragraph", "content": "{{date}}"},
                {"type": "paragraph", "content": "Dear {{recipient_name}},"},
                {"type": "paragraph", "content": "{{letter_content}}"},
                {"type": "paragraph", "content": "Sincerely,"},
                {"type": "paragraph", "content": "{{sender_name}}"},
            ]
        },
        is_public=True
    )
    
    # Add fields for letter template
    fields = [
        ("company_name", "Company Name", "text", True),
        ("date", "Date", "date", True),
        ("recipient_name", "Recipient Name", "text", True),
        ("letter_content", "Letter Content", "textarea", True),
        ("sender_name", "Sender Name", "text", True),
    ]
    
    for i, (name, label, field_type, required) in enumerate(fields):
        TemplateField.objects.create(
            template=letter_template,
            field_name=name,
            field_label=label,
            field_type=field_type,
            is_required=required,
            order=i
        )
    
    return letter_template
