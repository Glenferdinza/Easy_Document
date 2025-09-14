import os
import tempfile
from docx import Document
from docx2pdf import convert
import logging
import pythoncom

logger = logging.getLogger(__name__)


def convert_word_to_pdf(input_path, original_filename):
    """Convert Word document to PDF"""
    try:
        # Initialize COM
        pythoncom.CoInitialize()
        
        # Create output filename
        base_name = os.path.splitext(original_filename)[0]
        output_filename = f"{base_name}.pdf"
        
        # Create temporary output path
        output_dir = tempfile.mkdtemp()
        output_path = os.path.join(output_dir, output_filename)
        
        # Convert using docx2pdf
        convert(input_path, output_path)
        
        if not os.path.exists(output_path):
            raise Exception("PDF conversion failed - output file not created")
        
        return output_path
        
    except Exception as e:
        logger.error(f"Word to PDF conversion error: {str(e)}")
        raise Exception(f"Failed to convert Word to PDF: {str(e)}")
    finally:
        # Cleanup COM
        pythoncom.CoUninitialize()


def merge_word_documents(input_paths, output_format='docx'):
    """Merge multiple Word documents"""
    try:
        # Create merged document
        merged_doc = Document()
        
        # Add content from each document
        for i, input_path in enumerate(input_paths):
            if i > 0:
                # Add page break between documents
                merged_doc.add_page_break()
            
            # Load source document
            source_doc = Document(input_path)
            
            # Copy paragraphs
            for paragraph in source_doc.paragraphs:
                new_paragraph = merged_doc.add_paragraph()
                new_paragraph.text = paragraph.text
                new_paragraph.style = paragraph.style
            
            # Copy tables
            for table in source_doc.tables:
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.rows[i].cells[j].text = cell.text
        
        # Create output path
        output_dir = tempfile.mkdtemp()
        
        if output_format == 'docx':
            output_filename = "merged_document.docx"
            output_path = os.path.join(output_dir, output_filename)
            merged_doc.save(output_path)
        else:  # PDF
            # Initialize COM for PDF conversion
            pythoncom.CoInitialize()
            try:
                # Save as DOCX first, then convert to PDF
                temp_docx_path = os.path.join(output_dir, "temp_merged.docx")
                merged_doc.save(temp_docx_path)
                
                output_filename = "merged_document.pdf"
                output_path = os.path.join(output_dir, output_filename)
                convert(temp_docx_path, output_path)
            finally:
                pythoncom.CoUninitialize()
            
            # Clean up temporary DOCX
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
        
        if not os.path.exists(output_path):
            raise Exception(f"Merge failed - output file not created")
        
        return output_path
        
    except Exception as e:
        logger.error(f"Word merge error: {str(e)}")
        raise Exception(f"Failed to merge documents: {str(e)}")


def cleanup_file(file_path):
    """Safely remove a file"""
    try:
        if file_path and os.path.exists(file_path):
            os.unlink(file_path)
    except Exception as e:
        logger.warning(f"Failed to cleanup file {file_path}: {str(e)}")


def get_file_size_mb(file_path):
    """Get file size in MB"""
    try:
        size_bytes = os.path.getsize(file_path)
        return round(size_bytes / (1024 * 1024), 2)
    except:
        return 0
